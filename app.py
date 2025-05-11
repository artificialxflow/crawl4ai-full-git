import sys
import asyncio

if sys.platform.startswith("win"):
    asyncio.set_event_loop_policy(asyncio.WindowsProactorEventLoopPolicy())

import nest_asyncio
nest_asyncio.apply()

import streamlit as st
import requests
from bs4 import BeautifulSoup
from crawl4ai import AsyncWebCrawler, BrowserConfig, CrawlerRunConfig
from docx import Document
from fpdf import FPDF
import tempfile
import re
from dotenv import load_dotenv
import os
import json

# Load OpenRouter API key
load_dotenv('.env.local')
OPENROUTER_API_KEY = os.getenv('OPENROUTER_API_KEY')
OPENROUTER_MODEL = 'openai/gpt-4o'  # You can change this to your preferred model

# Inject custom CSS for RTL support and centering
st.markdown(
    """
    <style>
    body, .stApp {
        direction: rtl;
        text-align: right;
        font-family: Tahoma, Arial, sans-serif;
        background: #f9f9f9;
    }
    .centered-input {
        display: flex;
        justify-content: center;
        align-items: center;
        margin-bottom: 1.5rem;
    }
    .stTextInput > div > input {
        text-align: right;
    }
    .stDownloadButton > button {
        direction: rtl;
    }
    </style>
    """,
    unsafe_allow_html=True
)

st.title("🕸️ سایت‌مپ کرالر (RTL Sitemap Crawler)")

# State for parsed sitemap and crawl results
if 'sitemap_data' not in st.session_state:
    st.session_state['sitemap_data'] = None
if 'sitemap_error' not in st.session_state:
    st.session_state['sitemap_error'] = None
if 'crawl_results' not in st.session_state:
    st.session_state['crawl_results'] = None
if 'crawl_status' not in st.session_state:
    st.session_state['crawl_status'] = None
if 'is_crawling' not in st.session_state:
    st.session_state['is_crawling'] = False
if 'explanation' not in st.session_state:
    st.session_state['explanation'] = None
if 'knowledge_base' not in st.session_state:
    st.session_state['knowledge_base'] = None
if 'qa_answer' not in st.session_state:
    st.session_state['qa_answer'] = None
if 'archive' not in st.session_state:
    st.session_state['archive'] = []

# Sidebar: Archive of crawled sitemaps
st.sidebar.title('آرشیو سایت‌مپ‌ها')
if st.session_state['archive']:
    for i, entry in enumerate(st.session_state['archive']):
        st.sidebar.markdown(f"{i+1}. [{entry['url']}]({entry['url']})")
else:
    st.sidebar.info('هنوز هیچ سایت‌مپی خزیده نشده است.')

# Centered input for sitemap URL (empty by default)
with st.form(key="sitemap_form"):
    st.markdown('<div class="centered-input">', unsafe_allow_html=True)
    sitemap_url = st.text_input("آدرس سایت‌مپ را وارد کنید:", key="sitemap_url", value="")
    st.markdown('</div>', unsafe_allow_html=True)
    submit = st.form_submit_button("بررسی سایت‌مپ")

    if submit:
        st.session_state['sitemap_data'] = None
        st.session_state['sitemap_error'] = None
        st.session_state['crawl_results'] = None
        st.session_state['crawl_status'] = None
        st.session_state['is_crawling'] = False
        st.session_state['explanation'] = None
        st.session_state['knowledge_base'] = None
        st.session_state['qa_answer'] = None
        # Add to archive if not already present
        if sitemap_url and not any(entry['url'] == sitemap_url for entry in st.session_state['archive']):
            st.session_state['archive'].append({'url': sitemap_url})
        try:
            resp = requests.get(sitemap_url, timeout=10)
            resp.raise_for_status()
            soup = BeautifulSoup(resp.content, 'xml')
            urls = []
            for url in soup.find_all('url'):
                loc = url.find('loc').text if url.find('loc') else ''
                lastmod = url.find('lastmod').text if url.find('lastmod') else ''
                priority = url.find('priority').text if url.find('priority') else ''
                urls.append({'loc': loc, 'lastmod': lastmod, 'priority': priority})
            if urls:
                st.session_state['sitemap_data'] = urls
            else:
                st.session_state['sitemap_error'] = "هیچ URL معتبری در سایت‌مپ پیدا نشد."
        except Exception as e:
            st.session_state['sitemap_error'] = f"خطا در دریافت یا پردازش سایت‌مپ: {e}"

# Enable question input if sitemap is parsed and explanation is ready
question_disabled = not (st.session_state.get('sitemap_data') and not st.session_state.get('sitemap_error') and st.session_state.get('explanation'))

# Add a submit button for the question input
with st.form(key="qa_form"):
    st.markdown('<div class="centered-input">', unsafe_allow_html=True)
    question = st.text_input("سوال خود را درباره شرکت وارد کنید:", key="question", disabled=question_disabled)
    st.markdown('</div>', unsafe_allow_html=True)
    qa_submit = st.form_submit_button("پرسش سوال")

# Start crawling button
if st.session_state.get('sitemap_data'):
    if not st.session_state['is_crawling'] and not st.session_state.get('crawl_results'):
        if st.button('شروع خزیدن'):
            st.session_state['is_crawling'] = True
            st.rerun()

# Async crawl logic
async def crawl_all(urls):
    browser_conf = BrowserConfig(browser_type="firefox", headless=True, text_mode=True)
    crawl_status = {}
    crawl_results = {}
    async with AsyncWebCrawler(config=browser_conf) as crawler:
        results = await crawler.arun_many(
            urls=[u['loc'] for u in urls],
            config=CrawlerRunConfig()
        )
        for res in results:
            if res.success:
                crawl_status[res.url] = 'success'
                crawl_results[res.url] = res.markdown
            else:
                crawl_status[res.url] = 'error'
                crawl_results[res.url] = res.error_message
    return crawl_status, crawl_results

# Run crawling if triggered
if st.session_state.get('is_crawling') and st.session_state.get('sitemap_data') and not st.session_state.get('crawl_results'):
    with st.spinner('در حال خزیدن صفحات...'):
        loop = asyncio.get_event_loop()
        crawl_status, crawl_results = loop.run_until_complete(crawl_all(st.session_state['sitemap_data']))
        st.session_state['crawl_status'] = crawl_status
        st.session_state['crawl_results'] = crawl_results
        st.session_state['is_crawling'] = False
        st.rerun()

# --- LLM Content Analysis & Explanation Generation ---
def call_openrouter_llm(prompt, model=OPENROUTER_MODEL, max_tokens=1024, temperature=0.2):
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json"
    }
    data = {
        "model": model,
        "messages": [
            {"role": "system", "content": "شما یک تحلیل‌گر حرفه‌ای وب‌سایت هستید."},
            {"role": "user", "content": prompt}
        ],
        "max_tokens": max_tokens,
        "temperature": temperature
    }
    response = requests.post("https://openrouter.ai/api/v1/chat/completions", headers=headers, data=json.dumps(data))
    try:
        resp_json = response.json()
    except Exception:
        return f"[خطا در دریافت پاسخ از LLM: {response.text}]"
    if response.status_code == 200 and 'choices' in resp_json:
        return resp_json['choices'][0]['message']['content']
    elif 'error' in resp_json:
        return f"[خطا از LLM: {resp_json['error'].get('message', str(resp_json['error']))}]"
    else:
        return f"[خطا در دریافت پاسخ از LLM: {resp_json}]"

def build_analysis_prompt(all_text):
    return f"""
    با توجه به محتوای زیر که از صفحات مختلف یک شرکت استخراج شده است، لطفاً موارد زیر را به صورت دقیق و حرفه‌ای و با لحن انسانی و قابل استفاده برای بازاریابی و معرفی شرکت ارائه کن:
    1. خلاصه‌ای کامل و قابل فهم از شرکت و فعالیت‌های آن
    2. محصولات اصلی شرکت و توضیح هرکدام
    3. خدمات اصلی شرکت و توضیح هرکدام
    4. تخصص‌ها و مزیت‌های رقابتی شرکت
    5. کلمات کلیدی مهم مرتبط با محصولات و خدمات
    6. کلمات کلیدی مشتریان بالقوه
    7. یک ایمیل پیشنهادی برای ارسال به مشتریان بالقوه جهت معرفی شرکت و محصولات/خدمات
    8. هر نکته مهم دیگری که برای شناخت شرکت مفید است
    محتوای صفحات:
    ---
    {all_text}
    ---
    خروجی را به صورت منظم و قابل خواندن و با تیترهای مشخص ارائه کن."
    """

def build_qa_prompt(all_text, question):
    return f"""
    با توجه به اطلاعات زیر درباره شرکت، به سوال زیر به صورت دقیق و حرفه‌ای و با لحن انسانی پاسخ بده:
    اطلاعات شرکت:
    ---
    {all_text}
    ---
    سوال:
    {question}
    پاسخ را به فارسی و با جزئیات کافی ارائه کن.
    """

# Generate explanation after crawling using LLM
if st.session_state.get('crawl_results') and not st.session_state.get('explanation'):
    all_text = '\n\n'.join([v for v in st.session_state['crawl_results'].values() if isinstance(v, str) and len(v) > 30])
    MAX_CHARS = 40000
    if len(all_text) > MAX_CHARS:
        all_text = all_text[:MAX_CHARS]
    with st.spinner('در حال تحلیل محتوای سایت با هوش مصنوعی...'):
        prompt = build_analysis_prompt(all_text)
        explanation = call_openrouter_llm(prompt)
    st.session_state['explanation'] = explanation
    st.session_state['knowledge_base'] = all_text  # For QA

# Show explanation and export buttons
if st.session_state.get('explanation'):
    st.markdown('<div style="background:#fff;padding:1.5rem;border-radius:8px;box-shadow:0 2px 8px #eee;">', unsafe_allow_html=True)
    st.markdown(st.session_state['explanation'], unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

    # Export to Word
    doc = Document()
    doc.add_heading('توضیح شرکت و محصولات', 0)
    doc.add_paragraph(re.sub('<[^<]+?>', '', st.session_state['explanation']))
    word_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(word_file.name)
    with open(word_file.name, 'rb') as f:
        st.download_button('دانلود Word', f, file_name='company_explanation.docx')

    # Export to PDF
    pdf = FPDF()
    pdf.add_page()
    pdf.add_font('DejaVu', '', fname='C:/Windows/Fonts/arial.ttf', uni=True)  # Adjust font path as needed
    pdf.set_font('DejaVu', '', 14)
    for line in re.sub('<[^<]+?>', '', st.session_state['explanation']).split('\n'):
        pdf.cell(0, 10, line, ln=True, align='R')
    pdf_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
    pdf.output(pdf_file.name)
    with open(pdf_file.name, 'rb') as f:
        st.download_button('دانلود PDF', f, file_name='company_explanation.pdf')

# LLM-powered QA from knowledge base
if qa_submit and st.session_state.get('knowledge_base'):
    all_text = st.session_state['knowledge_base']
    MAX_CHARS = 40000
    if len(all_text) > MAX_CHARS:
        all_text = all_text[:MAX_CHARS]
    with st.spinner('در حال دریافت پاسخ از هوش مصنوعی...'):
        qa_prompt = build_qa_prompt(all_text, question)
        answer = call_openrouter_llm(qa_prompt)
    st.session_state['qa_answer'] = answer

if st.session_state.get('qa_answer'):
    st.info(st.session_state['qa_answer']) 